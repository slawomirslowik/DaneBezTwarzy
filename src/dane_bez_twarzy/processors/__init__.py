"""
Procesory dla różnych formatów plików.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from dane_bez_twarzy.core.anonymizer import Anonymizer


class FileProcessor(ABC):
    """Bazowa klasa dla procesorów plików."""
    
    @abstractmethod
    def process(
        self,
        input_path: Path,
        output_path: Path,
        anonymizer: 'Anonymizer',
        **kwargs: Any
    ) -> None:
        """
        Przetwarza plik.
        
        Args:
            input_path: Ścieżka do pliku wejściowego.
            output_path: Ścieżka do pliku wyjściowego.
            anonymizer: Instancja anonimizera.
            **kwargs: Dodatkowe argumenty.
        """
        pass


class TextProcessor(FileProcessor):
    """Processor dla plików tekstowych (.txt)."""
    
    def process(
        self,
        input_path: Path,
        output_path: Path,
        anonymizer: 'Anonymizer',
        **kwargs: Any
    ) -> None:
        """Przetwarza plik tekstowy."""
        # Wczytaj tekst
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Anonimizuj
        anonymized = anonymizer.anonymize_text(text)
        
        # Zapisz
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(anonymized)


class DocxProcessor(FileProcessor):
    """Processor dla plików Word (.docx)."""
    
    def process(
        self,
        input_path: Path,
        output_path: Path,
        anonymizer: 'Anonymizer',
        **kwargs: Any
    ) -> None:
        """Przetwarza plik DOCX."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Zainstaluj python-docx: pip install python-docx")
        
        # Wczytaj dokument
        doc = Document(input_path)
        
        # Anonimizuj paragrafy
        for paragraph in doc.paragraphs:
            if paragraph.text:
                anonymized = anonymizer.anonymize_text(paragraph.text)
                paragraph.text = anonymized
        
        # Anonimizuj tabele
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        anonymized = anonymizer.anonymize_text(cell.text)
                        cell.text = anonymized
        
        # Zapisz
        doc.save(output_path)


class ExcelProcessor(FileProcessor):
    """Processor dla plików Excel (.xlsx, .csv)."""
    
    def process(
        self,
        input_path: Path,
        output_path: Path,
        anonymizer: 'Anonymizer',
        **kwargs: Any
    ) -> None:
        """Przetwarza plik Excel/CSV."""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("Zainstaluj pandas: pip install pandas openpyxl")
        
        # Wczytaj plik
        if input_path.suffix == '.csv':
            df = pd.read_csv(input_path)
        else:
            df = pd.read_excel(input_path)
        
        # Anonimizuj każdą komórkę z tekstem
        for col in df.columns:
            if df[col].dtype == 'object':  # Kolumny tekstowe
                df[col] = df[col].apply(
                    lambda x: anonymizer.anonymize_text(str(x)) if pd.notna(x) else x
                )
        
        # Zapisz
        if output_path.suffix == '.csv':
            df.to_csv(output_path, index=False)
        else:
            df.to_excel(output_path, index=False)


class PDFProcessor(FileProcessor):
    """Processor dla plików PDF."""
    
    def process(
        self,
        input_path: Path,
        output_path: Path,
        anonymizer: 'Anonymizer',
        **kwargs: Any
    ) -> None:
        """Przetwarza plik PDF."""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("Zainstaluj PyPDF2: pip install PyPDF2")
        
        # Wczytaj PDF
        with open(input_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            writer = PyPDF2.PdfWriter()
            
            # Przetwórz każdą stronę
            for page in reader.pages:
                text = page.extract_text()
                
                if text:
                    anonymized = anonymizer.anonymize_text(text)
                    # Uwaga: PyPDF2 nie pozwala łatwo modyfikować tekstu w PDF
                    # Ta implementacja jest uproszczona
                    # W produkcji użyj np. reportlab do generowania nowego PDF
                    pass
                
                writer.add_page(page)
            
            # Zapisz (bez modyfikacji tekstu - wymaga bardziej zaawansowanej implementacji)
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
        
        # Uwaga: Ta implementacja tylko kopiuje PDF
        # Pełna anonimizacja PDF wymaga bardziej zaawansowanych narzędzi


def get_processor(file_extension: str) -> FileProcessor:
    """
    Zwraca odpowiedni processor dla typu pliku.
    
    Args:
        file_extension: Rozszerzenie pliku (z kropką, np. '.txt').
        
    Returns:
        Instancja processora.
    """
    processors = {
        '.txt': TextProcessor,
        '.docx': DocxProcessor,
        '.xlsx': ExcelProcessor,
        '.csv': ExcelProcessor,
        '.pdf': PDFProcessor,
    }
    
    processor_class = processors.get(file_extension.lower())
    if not processor_class:
        raise ValueError(f"Nieobsługiwany format pliku: {file_extension}")
    
    return processor_class()
